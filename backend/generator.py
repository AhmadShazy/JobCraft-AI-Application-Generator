import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_margins(doc, margin_in_inches=1.0):
    """Sets standard margins on all sections of the document."""
    for section in doc.sections:
        section.top_margin = Inches(margin_in_inches)
        section.bottom_margin = Inches(margin_in_inches)
        section.left_margin = Inches(margin_in_inches)
        section.right_margin = Inches(margin_in_inches)

def add_bottom_border(paragraph):
    """Adds a clean bottom border line to a heading paragraph (classic resume divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # 6/8 pt size
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '000000')  # Solid Black
    pBdr.append(bottom)
    pPr.append(pBdr)

def apply_font_settings(run, name="Arial", size_pt=10, color_rgb=(0, 0, 0), bold=False, italic=False):
    """Applies font styles to a run."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic

def add_markdown_runs(paragraph, text, size_pt=10, default_bold=False, default_italic=False, color_rgb=(0, 0, 0)):
    """
    Parses simple markdown bold (**) and italic (*) and adds them as formatted runs to the paragraph.
    """
    # Split text on **bold** or *italic* tags
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            apply_font_settings(run, size_pt=size_pt, bold=True, italic=default_italic, color_rgb=color_rgb)
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            apply_font_settings(run, size_pt=size_pt, bold=default_bold, italic=True, color_rgb=color_rgb)
        else:
            run = paragraph.add_run(token)
            apply_font_settings(run, size_pt=size_pt, bold=default_bold, italic=default_italic, color_rgb=color_rgb)

def generate_resume_docx(resume_text: str, profile_data: dict, output_path: str):
    """
    Generates an ATS-friendly tailored resume in .docx format by parsing
    the plain text resume_text using regex section markers.
    Uses Arial, black-only color scheme, standard 1-inch margins,
    and a clean single-column structure.
    """
    doc = Document()
    set_margins(doc, 1.0)
    
    # Page printable width is 8.5" (letter width) - 2" (margins) = 6.5"
    # Set up a right tab-stop at 6.5" for aligning dates/durations
    right_tab_stop_position = Inches(6.5)

    # Parse sections using regex
    parts = re.split(r'===\s*([^=]+?)\s*===', resume_text)
    sections = {}
    for i in range(1, len(parts), 2):
        section_name = parts[i].strip().upper()
        section_content = parts[i+1].strip()
        sections[section_name] = section_content

    # Print statement showing successfully parsed sections
    print(f"Successfully parsed sections: {list(sections.keys())}")

    # 1. Header (Name, Dynamic Tagline, and Contact Info)
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(profile_data.get("name", "Ahmad Sheraz"))
    apply_font_settings(name_run, size_pt=16, bold=True)

    # Center-aligned Dynamic Tagline under candidate name
    tagline_val = sections.get("TAGLINE")
    if tagline_val:
        tagline_p = doc.add_paragraph()
        tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline_p.paragraph_format.space_after = Pt(2)
        tagline_run = tagline_p.add_run(tagline_val.strip())
        apply_font_settings(tagline_run, size_pt=11, bold=True, italic=True)

    # Compile contact details dynamically
    email = profile_data.get("email", "")
    phone = profile_data.get("phone", "")
    location = profile_data.get("location", "")
    linkedin = profile_data.get("linkedin", "")
    github = profile_data.get("github", "")
    portfolio = profile_data.get("portfolio", "")
    
    contact_parts = [email, phone, location]
    if linkedin:
        contact_parts.append(linkedin)
    if github:
        contact_parts.append(github)
    if portfolio:
        contact_parts.append(portfolio)
    
    contact_text = " | ".join([p for p in contact_parts if p])
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(12)
    contact_run = contact_p.add_run(contact_text)
    apply_font_settings(contact_run, size_pt=9.5)

    # 2. Iterate through sections sequentially in the order they were parsed
    for i in range(1, len(parts), 2):
        section_name = parts[i].strip().upper()
        section_content = parts[i+1].strip()
        
        if not section_content:
            continue
            
        if section_name == "TAGLINE":
            # Tagline has already been rendered in the header
            continue
            
        elif "SUMMARY" in section_name:
            summary_title_p = doc.add_paragraph()
            summary_title_p.paragraph_format.space_before = Pt(8)
            summary_title_p.paragraph_format.space_after = Pt(4)
            add_bottom_border(summary_title_p)
            title_run = summary_title_p.add_run("PROFESSIONAL SUMMARY")
            apply_font_settings(title_run, size_pt=11.5, bold=True)

            summary_p = doc.add_paragraph()
            summary_p.paragraph_format.space_after = Pt(8)
            summary_p.paragraph_format.line_spacing = 1.15
            add_markdown_runs(summary_p, section_content, size_pt=10)

        elif "SKILL" in section_name:
            skills_title_p = doc.add_paragraph()
            skills_title_p.paragraph_format.space_before = Pt(8)
            skills_title_p.paragraph_format.space_after = Pt(4)
            add_bottom_border(skills_title_p)
            title_run = skills_title_p.add_run("TECHNICAL SKILLS")
            apply_font_settings(title_run, size_pt=11.5, bold=True)

            lines = section_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if ":" in line:
                    cat, items_str = [p.strip() for p in line.split(":", 1)]
                    
                    skill_p = doc.add_paragraph()
                    skill_p.paragraph_format.space_after = Pt(3)
                    skill_p.paragraph_format.line_spacing = 1.1
                    
                    cat_run = skill_p.add_run(f"{cat}: ")
                    apply_font_settings(cat_run, size_pt=10, bold=True)
                    
                    list_run = skill_p.add_run(items_str)
                    apply_font_settings(list_run, size_pt=10)

        elif "EXPERIENCE" in section_name:
            exp_title_p = doc.add_paragraph()
            exp_title_p.paragraph_format.space_before = Pt(10)
            exp_title_p.paragraph_format.space_after = Pt(4)
            add_bottom_border(exp_title_p)
            title_run = exp_title_p.add_run("PROFESSIONAL EXPERIENCE")
            apply_font_settings(title_run, size_pt=11.5, bold=True)

            lines = section_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("-") or line.startswith("*"):
                    bullet = line.lstrip("-* ").strip()
                    bullet_p = doc.add_paragraph(style='List Bullet')
                    bullet_p.paragraph_format.space_after = Pt(2)
                    bullet_p.paragraph_format.line_spacing = 1.1
                    add_markdown_runs(bullet_p, bullet, size_pt=10)
                elif "|" in line:
                    job_parts = [p.strip() for p in line.split("|")]
                    title = job_parts[0] if len(job_parts) > 0 else ""
                    company = job_parts[1] if len(job_parts) > 1 else ""
                    location = job_parts[2] if len(job_parts) > 2 else ""
                    duration = job_parts[3] if len(job_parts) > 3 else ""

                    job_p = doc.add_paragraph()
                    job_p.paragraph_format.tab_stops.add_tab_stop(right_tab_stop_position, WD_TAB_ALIGNMENT.RIGHT)
                    job_p.paragraph_format.space_before = Pt(4)
                    job_p.paragraph_format.space_after = Pt(1)

                    title_company = f"{title} – {company}"
                    title_run = job_p.add_run(title_company)
                    apply_font_settings(title_run, size_pt=10.5, bold=True)

                    duration_run = job_p.add_run(f"\t{duration}")
                    apply_font_settings(duration_run, size_pt=10, bold=True)

                    if location:
                        loc_p = doc.add_paragraph()
                        loc_p.paragraph_format.space_after = Pt(2)
                        loc_run = loc_p.add_run(location)
                        apply_font_settings(loc_run, size_pt=9.5, italic=True)

        elif "PROJECT" in section_name:
            proj_title_p = doc.add_paragraph()
            proj_title_p.paragraph_format.space_before = Pt(10)
            proj_title_p.paragraph_format.space_after = Pt(4)
            add_bottom_border(proj_title_p)
            title_run = proj_title_p.add_run("PROJECTS")
            apply_font_settings(title_run, size_pt=11.5, bold=True)

            lines = section_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("-") or line.startswith("*"):
                    bullet = line.lstrip("-* ").strip()
                    bullet_p = doc.add_paragraph(style='List Bullet')
                    bullet_p.paragraph_format.space_after = Pt(2)
                    bullet_p.paragraph_format.line_spacing = 1.1
                    add_markdown_runs(bullet_p, bullet, size_pt=10)
                elif "|" in line:
                    proj_parts = [p.strip() for p in line.split("|")]
                    name = proj_parts[0] if len(proj_parts) > 0 else ""
                    duration = proj_parts[1] if len(proj_parts) > 1 else ""
                    stack = proj_parts[2] if len(proj_parts) > 2 else ""

                    proj_p = doc.add_paragraph()
                    proj_p.paragraph_format.tab_stops.add_tab_stop(right_tab_stop_position, WD_TAB_ALIGNMENT.RIGHT)
                    proj_p.paragraph_format.space_before = Pt(4)
                    proj_p.paragraph_format.space_after = Pt(1)

                    name_run = proj_p.add_run(name)
                    apply_font_settings(name_run, size_pt=10.5, bold=True)

                    duration_run = proj_p.add_run(f"\t{duration}")
                    apply_font_settings(duration_run, size_pt=10, bold=True)

                    if stack:
                        stack_p = doc.add_paragraph()
                        stack_p.paragraph_format.space_after = Pt(2)
                        stack_run = stack_p.add_run(f"Technologies: {stack}")
                        apply_font_settings(stack_run, size_pt=9.5, italic=True)

        elif "EDUCATION" in section_name:
            edu_title_p = doc.add_paragraph()
            edu_title_p.paragraph_format.space_before = Pt(10)
            edu_title_p.paragraph_format.space_after = Pt(4)
            add_bottom_border(edu_title_p)
            title_run = edu_title_p.add_run("EDUCATION")
            apply_font_settings(title_run, size_pt=11.5, bold=True)

            lines = section_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if "|" in line:
                    edu_parts = [p.strip() for p in line.split("|")]
                    degree = edu_parts[0] if len(edu_parts) > 0 else ""
                    institution = edu_parts[1] if len(edu_parts) > 1 else ""
                    duration = edu_parts[2] if len(edu_parts) > 2 else ""
                    note = edu_parts[3] if len(edu_parts) > 3 else ""

                    edu_p = doc.add_paragraph()
                    edu_p.paragraph_format.tab_stops.add_tab_stop(right_tab_stop_position, WD_TAB_ALIGNMENT.RIGHT)
                    edu_p.paragraph_format.space_before = Pt(3)
                    edu_p.paragraph_format.space_after = Pt(1)

                    edu_run = edu_p.add_run(degree)
                    apply_font_settings(edu_run, size_pt=10.5, bold=True)

                    duration_run = edu_p.add_run(f"\t{duration}")
                    apply_font_settings(duration_run, size_pt=10, bold=True)

                    inst_p = doc.add_paragraph()
                    inst_p.paragraph_format.space_after = Pt(2)
                    inst_text = institution
                    if note:
                        inst_text += f" ({note})"
                    inst_run = inst_p.add_run(inst_text)
                    apply_font_settings(inst_run, size_pt=10, italic=True)

        elif "CERTIFICATION" in section_name:
            cert_title_p = doc.add_paragraph()
            cert_title_p.paragraph_format.space_before = Pt(10)
            cert_title_p.paragraph_format.space_after = Pt(4)
            add_bottom_border(cert_title_p)
            title_run = cert_title_p.add_run("CERTIFICATIONS")
            apply_font_settings(title_run, size_pt=11.5, bold=True)

            lines = section_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                cert = line.lstrip("-* ").strip()
                if "|" in cert:
                    cert_parts = [p.strip() for p in cert.split("|")]
                    name = cert_parts[0] if len(cert_parts) > 0 else ""
                    issuer = cert_parts[1] if len(cert_parts) > 1 else ""
                    date = cert_parts[2] if len(cert_parts) > 2 else ""
                    text = f"{name} – {issuer} ({date})"
                else:
                    text = cert

                cert_p = doc.add_paragraph(style='List Bullet')
                cert_p.paragraph_format.space_after = Pt(2)
                cert_run = cert_p.add_run(text)
                apply_font_settings(cert_run, size_pt=10)

        elif "VOLUNTEER" in section_name:
            vol_title_p = doc.add_paragraph()
            vol_title_p.paragraph_format.space_before = Pt(10)
            vol_title_p.paragraph_format.space_after = Pt(4)
            add_bottom_border(vol_title_p)
            title_run = vol_title_p.add_run("VOLUNTEER EXPERIENCE")
            apply_font_settings(title_run, size_pt=11.5, bold=True)

            lines = section_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("-") or line.startswith("*"):
                    bullet = line.lstrip("-* ").strip()
                    bullet_p = doc.add_paragraph(style='List Bullet')
                    bullet_p.paragraph_format.space_after = Pt(2)
                    bullet_p.paragraph_format.line_spacing = 1.1
                    add_markdown_runs(bullet_p, bullet, size_pt=10)
                elif "|" in line:
                    vol_parts = [p.strip() for p in line.split("|")]
                    role = vol_parts[0] if len(vol_parts) > 0 else ""
                    org = vol_parts[1] if len(vol_parts) > 1 else ""
                    duration = vol_parts[2] if len(vol_parts) > 2 else ""

                    vol_p = doc.add_paragraph()
                    vol_p.paragraph_format.tab_stops.add_tab_stop(right_tab_stop_position, WD_TAB_ALIGNMENT.RIGHT)
                    vol_p.paragraph_format.space_before = Pt(3)
                    vol_p.paragraph_format.space_after = Pt(1)

                    role_org = f"{role} – {org}"
                    role_run = vol_p.add_run(role_org)
                    apply_font_settings(role_run, size_pt=10.5, bold=True)

                    duration_run = vol_p.add_run(f"\t{duration}")
                    apply_font_settings(duration_run, size_pt=10, bold=True)

        elif "LANGUAGE" in section_name:
            lang_title_p = doc.add_paragraph()
            lang_title_p.paragraph_format.space_before = Pt(10)
            lang_title_p.paragraph_format.space_after = Pt(4)
            add_bottom_border(lang_title_p)
            title_run = lang_title_p.add_run("LANGUAGES")
            apply_font_settings(title_run, size_pt=11.5, bold=True)

            lines = section_content.split("\n")
            lang_strings = []
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                lang_item = line.lstrip("-* ").strip()
                if "|" in lang_item:
                    lang_parts = [p.strip() for p in lang_item.split("|")]
                    lang_strings.append(f"{lang_parts[0]} ({lang_parts[1]})")
                else:
                    lang_strings.append(lang_item)

            lang_p = doc.add_paragraph()
            lang_p.paragraph_format.space_after = Pt(8)
            lang_run = lang_p.add_run(", ".join(lang_strings))
            apply_font_settings(lang_run, size_pt=10)

    # Make parent directories if they don't exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def generate_cover_letter_docx(cl_data: dict, name: str, email: str, phone: str, links: list, output_path: str):
    """
    Generates a professional cover letter in .docx format.
    Uses Arial, black-only color scheme, standard 1-inch margins.
    Formatted as block paragraphs (no indent, with spaces between paragraphs).
    """
    doc = Document()
    set_margins(doc, 1.0)

    # 1. Header (Left-aligned)
    header_p = doc.add_paragraph()
    header_p.paragraph_format.space_after = Pt(2)
    name_run = header_p.add_run(name)
    apply_font_settings(name_run, size_pt=12, bold=True)

    info_p = doc.add_paragraph()
    info_p.paragraph_format.space_after = Pt(18)
    contact_info = f"Email: {email} | Phone: {phone}"
    if links:
        contact_info += " | " + " | ".join(links)
    info_run = info_p.add_run(contact_info)
    apply_font_settings(info_run, size_pt=9.5)

    # 2. Date
    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_after = Pt(12)
    from datetime import datetime
    now = datetime.now()
    default_date_str = f"{now.strftime('%B')} {now.day}, {now.year}"
    date_run = date_p.add_run(cl_data.get("date") or default_date_str)
    apply_font_settings(date_run, size_pt=10)

    # 3. Recipient
    rec_p = doc.add_paragraph()
    rec_p.paragraph_format.space_after = Pt(12)
    rec_run = rec_p.add_run(f"To,\nThe Hiring Team\n{cl_data.get('recipient_company', 'Target Company')}")
    apply_font_settings(rec_run, size_pt=10)

    # 4. Subject line
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run(cl_data.get("subject", "RE: Application for Position"))
    apply_font_settings(sub_run, size_pt=10, bold=True)

    # 5. Salutation
    sal_p = doc.add_paragraph()
    sal_p.paragraph_format.space_after = Pt(12)
    sal_run = sal_p.add_run(cl_data.get("salutation", "Dear Hiring Team,"))
    apply_font_settings(sal_run, size_pt=10)

    # 6. Body Paragraphs
    for para_text in cl_data.get("paragraphs", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
        add_markdown_runs(p, para_text, size_pt=10)

    # 7. Sign-off
    sign_p = doc.add_paragraph()
    sign_p.paragraph_format.space_before = Pt(12)
    sign_run = sign_p.add_run(cl_data.get("sign_off", f"Sincerely,\n\n{name}"))
    apply_font_settings(sign_run, size_pt=10)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
